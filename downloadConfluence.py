import config
import shutil
import requests
from requests.auth import HTTPBasicAuth
import xml.etree.ElementTree as ET
import re
import docx
import os
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from bs4 import BeautifulSoup
import warnings
# warnings.filterwarnings("ignore", category=UserWarning, module='docx')
import urllib.parse
import logging


def convert_html_to_docx(file_path):
    # Leggi il file HTML
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
    except FileNotFoundError:
        print(f"Error: File {file_path} not found.")
        return

    # Crea un nuovo documento DOCX
    doc = Document()

    # Utilizza BeautifulSoup per parsare il contenuto HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Aggiungi il contenuto HTML al documento DOCX
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'table']):
        if element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'img':
                    # Aggiungi l'immagine al paragrafo
                    img_path = os.path.join(os.path.dirname(file_path), child['src'])
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6))
                else:
                    # Aggiungi il testo al paragrafo
                    run = p.add_run(child.text)
        elif element.name == 'ul':
            # Aggiungi una nuova lista non ordinata al documento
            for child in element.children:
                if child.name == 'li':
                    p = doc.add_paragraph()
                    run = p.add_run('• ')
                    run = p.add_run(child.text)
        elif element.name == 'ol':
            # Aggiungi una nuova lista ordinata al documento
            for i, child in enumerate(element.children):
                if child.name == 'li':
                    p = doc.add_paragraph()
                    run = p.add_run(f'{i + 1}. ')
                    run = p.add_run(child.text)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            doc.add_heading(element.text, level=int(element.name[1]))
        elif element.name == 'table':
            # Aggiungi una nuova tabella al documento
            rows = element.find_all('tr')
            cols = max(len(row.find_all(['th', 'td'])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols, style='TableGrid')
            for i, row in enumerate(rows):
                for j, cell in enumerate(row.find_all(['th', 'td'])):
                    if j < cols:
                        if cell.name == 'th':
                            # Aggiungi l'intestazione della colonna
                            table.cell(i, j).text = cell.text
                            table.cell(i, j).paragraphs[0].runs[0].font.bold = True
                        else:
                            table.cell(i, j).text = cell.text

    # Salva il file DOCX
    file_name, file_extension = os.path.splitext(file_path)
    docx_file = file_name + '_confluence.docx'
    try:
        doc.save(docx_file)
    except Exception as e:
        print(f"Error: {e}")


def convert_txt_html(file_path):
    # Verifica se il file esiste
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found.")
        return None

    # Ottieni il percorso del file e il nome del file senza estensione
    file_dir, file_name = os.path.split(file_path)
    file_name_without_ext = os.path.splitext(file_name)[0]

    # Rinomina il file da .txt a .html
    new_file_name = f"{file_name_without_ext}.html"
    new_file_path = os.path.join(file_dir, new_file_name)

    # Copia il file
    shutil.copy(file_path, new_file_path)

    return new_file_path


def convert_image_inclusions(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    """
        <ac:image ac:height="250"><ri:attachment ri:filename="image2017-4-26_15-14-0.png" /></ac:image>
    
        <ac:image ac:height="250">
        <ri:attachment ri:filename="image2017-4-26_15-14-0.png"/>
        </ac:image>
    """
    pattern = r'<ac:image ac:height="(\d+)"><ri:attachment ri:filename="([^"]+)" /></ac:image>'
    replacement = r'<img src="\2" height="\1">'

    new_content = re.sub(pattern, replacement, content, flags=re.DOTALL)

    with open(file_path, 'w', encoding='utf-8') as new_file:
        new_file.write(new_content)

def convert_blocks(file_path):
    """
    Converte i blocchi 'info', 'warning' e 'code' in una tabella formattata.
    """

    # Configurazione base di logging per tracciare i messaggi di debug
    logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(message)s')
    logging.debug("Inizio elaborazione del file: %s", file_path)

    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
            logging.debug("Contenuto del file letto con successo.")
    except Exception as e:
        logging.error("Errore durante l'apertura del file: %s", e)
        return

    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        logging.debug("Contenuto HTML analizzato correttamente.")
    except Exception as e:
        logging.error("Errore durante l'analisi del contenuto HTML: %s", e)
        return

    # Elaborazione dei blocchi 'info'
    try:
        for info_block in soup.find_all("ac:structured-macro", {"ac:name": "info"}):
            logging.debug("Trovato blocco 'info'")

            info_tag = info_block.find("ac:rich-text-body")
            if info_tag:
                info_content = info_tag.get_text(separator="\n", strip=False)
                logging.debug("Contenuto info estratto: %s", info_content)

                # Creazione della tabella per 'info'
                table_info = soup.new_tag("table")
                tr1_info = soup.new_tag("tr")
                td1_info = soup.new_tag("td")
                pre1_info = soup.new_tag("pre")
                pre1_info.string = "--**-- Info"
                td1_info.append(pre1_info)
                tr1_info.append(td1_info)

                tr_info = soup.new_tag("tr")
                td_info = soup.new_tag("td")
                pre_info = soup.new_tag("pre")
                pre_info.string = info_content
                td_info.append(pre_info)
                tr_info.append(td_info)

                table_info.append(tr1_info)
                table_info.append(tr_info)

                # Sostituzione del blocco 'info' con la tabella
                info_block.replace_with(table_info)
                logging.debug("Blocco 'info' sostituito con la tabella.")
            else:
                logging.warning("Blocco 'info' senza contenuto valido.")
    except Exception as e:
        logging.error("Errore durante l'elaborazione dei blocchi 'info': %s", e)

    # Elaborazione dei blocchi 'warning'
    try:
        for warning_block in soup.find_all("ac:structured-macro", {"ac:name": "warning"}):
            logging.debug("Trovato blocco 'warning'")

            warning_tag = warning_block.find("ac:rich-text-body")
            if warning_tag:
                warning_content =  warning_tag.get_text(separator="\n", strip=False)
                logging.debug("Contenuto warning estratto: %s", warning_content)

                # Creazione della tabella per 'warning'
                table_warning = soup.new_tag("table")
                tr1_warning = soup.new_tag("tr")
                td1_warning = soup.new_tag("td")
                pre1_warning = soup.new_tag("pre")
                pre1_warning.string = "--**-- Warning"
                td1_warning.append(pre1_warning)
                tr1_warning.append(td1_warning)

                tr_warning = soup.new_tag("tr")
                td_warning = soup.new_tag("td")
                pre_warning = soup.new_tag("pre")
                pre_warning.string = warning_content
                td_warning.append(pre_warning)
                tr_warning.append(td_warning)

                table_warning.append(tr1_warning)
                table_warning.append(tr_warning)

                # Sostituzione del blocco 'warning' con la tabella
                warning_block.replace_with(table_warning)
                logging.debug("Blocco 'warning' sostituito con la tabella.")
            else:
                logging.warning("Blocco 'warning' senza contenuto valido.")
    except Exception as e:
        logging.error("Errore durante l'elaborazione dei blocchi 'warning': %s", e)

    # Elaborazione dei blocchi 'code'
    try:
        for code_block in soup.find_all("ac:structured-macro", {"ac:name": "code"}):
            logging.debug("Trovato blocco 'code'")

            cdata_tag = code_block.find("ac:plain-text-body")
            if cdata_tag and cdata_tag.string:
                code_content = cdata_tag.string.strip()
                if code_content.startswith("<![CDATA[") and code_content.endswith("]]>"):
                    code_content = code_content[9:-3].strip()

                logging.debug("Contenuto codice estratto: %s", code_content)

                # Creazione della tabella per 'code'
                table = soup.new_tag("table")
                tr = soup.new_tag("tr")
                td = soup.new_tag("td")
                pre = soup.new_tag("pre")
                pre.string = code_content
                td.append(pre)
                tr.append(td)
                table.append(tr)

                # Sostituzione del blocco 'code' con la tabella
                code_block.replace_with(table)
                logging.debug("Blocco 'code' sostituito con la tabella.")
            else:
                logging.warning("Blocco 'code' senza contenuto valido.")
    except Exception as e:
        logging.error("Errore durante l'elaborazione dei blocchi 'code': %s", e)

    # Scrittura del nuovo contenuto nel file
    try:
        with open(file_path, 'w', encoding='utf-8') as new_file:
            new_file.write(str(soup))
            logging.debug("File salvato correttamente.")
    except Exception as e:
        logging.error("Errore durante il salvataggio del file: %s", e)



def convert_attachment_inclusions(file_path):
    """
    <ac:structured-macro ac:macro-id="f784c3c2-8991-4dd2-9a2a-db23ace74e10" ac:name="view-file" ac:schema-version="1">
        <ac:parameter ac:name="name">
            <ri:attachment ri:filename="local_webdav\01_Clienti\_CLIENTI\CN\Analisi PPN.docx"></ri:attachment>
        </ac:parameter>
        <ac:parameter ac:name="height">250</ac:parameter>
    </ac:structured-macro>
    <ac:structured-macro ac:macro-id="d0dd93be-1cb3-4501-b973-ace9c30ff01e" ac:name="view-file" ac:schema-version="1">
        <ac:parameter ac:name="name">
            <ri:attachment ri:filename="local_webdav\01_Clienti\_CLIENTI\CN\Manuale__PPN.docx"></ri:attachment>
        </ac:parameter>
        <ac:parameter ac:name="height">250</ac:parameter>
    </ac:structured-macro>
    <ac:structured-macro ac:macro-id="e56a02b5-b741-41c0-8a30-cd8552797a2e" ac:name="view-file" ac:schema-version="1">
        <ac:parameter ac:name="name">
            <ri:attachment
                    ri:filename="local_webdav\01_Clienti\_CLIENTI\CN\Manuale__PPN_Freschissimo.docx"></ri:attachment>
        </ac:parameter>
        <ac:parameter ac:name="height">250</ac:parameter>
    </ac:structured-macro>
    :param file_path:
    :return:
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    attachments = soup.find_all('ac:structured-macro')

    for attachment in attachments:
        name_element = attachment.find('ac:parameter', {'ac:name': 'name'})
        if name_element is not None:
            filename = name_element.find('ri:attachment')['ri:filename']
            filename = filename.split('\\')[-1]  # prendi solo il nome del file

            # Sostituisci il testo all'interno dell'elemento 'ac:structured-macro' con solo il nome del file
            attachment.replace_with(filename + '\n')

    with open(file_path, 'w', encoding='utf-8') as new_file:
        new_file.write(str(soup))


def pulisci_percorso(percorso):
    """
    Pulisce il percorso di cartelle e file per evitare errori di creazione.

    :param percorso: il percorso da pulire
    :return: il percorso pulito
    """
    percorso = re.sub(r'[*?:<>|"]', '_', percorso)
    percorso = re.sub(r'\\', '_', percorso)
    percorso = re.sub(r':', '_', percorso)
    percorso = re.sub(r'%20', ' ', percorso)
    percorso = re.sub(r'%40', '_', percorso)
    percorso = re.sub(r'%3d', '=', percorso)
    percorso = re.sub(r'%3a', '=', percorso)
    percorso = urllib.parse.unquote(percorso)

    percorso = os.path.normpath(percorso)
    return percorso


# Funzione per scaricare un file e salvarlo nella struttura di cartelle
def download_webdav_file(file_url, base_url, username, password, local_root="local_webdav"):
    try:
        # Assicurati che il file_url inizi con base_url
        if not file_url.startswith(base_url):
            print("Il file non appartiene al percorso specificato.")
            return

        # Estrai il percorso relativo dalla parte specificata
        relative_path = pulisci_percorso(file_url[len(base_url):])

        # Crea il percorso locale replicando la struttura originale
        local_path = os.path.join(local_root, relative_path)

        # Crea le directory necessarie
        os.makedirs(os.path.dirname(local_path), exist_ok=True)

        if not os.path.exists(local_path):
            # Scarica il file
            response = requests.get(file_url, auth=HTTPBasicAuth(username, password), stream=True)
            if response.status_code == 200:
                # Salva il file localmente
                with open(local_path, "wb") as file:
                    for chunk in response.iter_content(chunk_size=8192):
                        file.write(chunk)
                print(f"File salvato con successo: {local_path}")
                cartella_precedente = os.path.basename(os.path.dirname(local_path))
                nome_file = os.path.basename(local_path)
                nome_file_senza_estensione = os.path.splitext(nome_file)[0]
                nome_file_ext = os.path.splitext(nome_file)[1]
                print(f"cartella_precedente: {cartella_precedente}")
                print(f"nome_file: {nome_file}")
                print(f"nome_file_senza_estensione: {nome_file_senza_estensione}")
                print(f"nome_file_ext: {nome_file_ext}")
                if cartella_precedente == nome_file_senza_estensione and nome_file_ext == ".txt":
                    newhtml = convert_txt_html(local_path)
                    convert_blocks(newhtml)
                    convert_image_inclusions(newhtml)
                    convert_attachment_inclusions(newhtml)
                    convert_html_to_docx(newhtml)
                    print(f"File txt convertito in html: {local_path}")
            else:
                print(f"Errore durante il download: {response.status_code} - {response.text}")
            print("File scaricato con successo!")
        else:
            print("File già esistente, non scaricato.")
    except Exception as e:
        print(f"Errore: {e}")


# Funzione per elencare il contenuto della cartella WebDAV
def list_webdav_content(url, username, password):
    print('-------------')
    print(url)

    headers = {
        "Depth": "1"  # Richiede solo i contenuti della directory principale
    }
    # Costruisce la richiesta PROPFIND
    response = requests.request(
        "PROPFIND", url, auth=HTTPBasicAuth(username, password), headers=headers
    )

    if response.status_code == 207:  # Multi-Status indica una risposta valida
        try:
            # Analizza il contenuto XML
            root = ET.fromstring(response.content)
            namespace = {"d": "DAV:"}  # Namespace DAV

            # Trova tutti i nodi con informazioni sui file e cartelle
            items = root.findall(".//d:response", namespace)
            files = []
            folders = []

            for item in items:
                href = item.find("d:href", namespace)
                if href is not None:
                    path = href.text
                    # Distingui file e cartelle basandoti sullo slash finale
                    if path.endswith("/"):
                        folders.append(path)
                    else:
                        files.append(path)

            # Stampa l'elenco di cartelle
            print("Cartelle:")
            for folder in folders:
                print(folder)
                if len(folders) > 0 and url != folder:  # se la cartella non è vuota
                    list_webdav_content(folder, config.USERNAME, config.PASSWORD)

            # Stampa l'elenco di file
            print("\nFile:")
            for file in files:
                print(file)
                # Esempio di utilizzo
                # file_url = "https://docs.ditechonline.it/plugins/servlet/confluence/default/Global/PCD/Promo%20Collaboration%20Development/01_Clienti/example/file.txt"
                download_webdav_file(file, config.BASE_WEBDAV_URL, config.USERNAME, config.PASSWORD)

        except ET.ParseError as e:
            print(f"Errore nell'analisi della risposta XML: {e}")
    else:
        print(f"Errore: {response.status_code} - {response.text}")


# Esegui la funzione
list_webdav_content(config.BASE_WEBDAV_URL, config.USERNAME, config.PASSWORD)
