import os
import re
import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup
from html2docx import html2docx
from docx import Document

# Configurazione
BASE_URL =
API_USER =
API_PASSWORD =
SPACE_KEY = "PCD"
OUTPUT_DIR = "./confluence_export"

# Creazione directory di output
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Funzione per pulire i nomi delle directory e dei file
# Funzione per pulire i nomi delle directory e dei file
def clean_name(name):
    # Sostituzione specifica
    name = name.replace("Promo Collaboration Development", "PCD")
    
    # Sostituzione dei caratteri non validi nei nomi dei file e delle cartelle
    return re.sub(r'[<>:"/\\|?*]', '-', name)

# Funzione per ottenere la lista delle pagine nello spazio con dettagli sugli antenati
def get_pages_with_hierarchy(space_key):
    url = f"{BASE_URL}/rest/api/content"
    params = {
        "spaceKey": space_key,
        "expand": "ancestors,body.storage",
        "limit": 100
    }
    response = requests.get(url, auth=HTTPBasicAuth(API_USER, API_PASSWORD), params=params)
    response.raise_for_status()
    return response.json()["results"]

# Funzione per ottenere il percorso della gerarchia
def get_hierarchy_path(page):
    ancestors = page.get("ancestors", [])
    path = [clean_name(ancestor["title"]) for ancestor in ancestors]
    path.append(clean_name(page["title"]))
    return os.path.join(*path)

# Funzione per scaricare gli allegati di una pagina
def download_attachments(page_id, page_path):
    url = f"{BASE_URL}/rest/api/content/{page_id}/child/attachment"
    response = requests.get(url, auth=HTTPBasicAuth(API_USER, API_PASSWORD))
    response.raise_for_status()
    attachments = response.json()["results"]

    attachments_dir = os.path.join(page_path, "attachments")
    os.makedirs(attachments_dir, exist_ok=True)

    for attachment in attachments:
        attachment_url = attachment["_links"]["download"]
        file_name = clean_name(attachment["title"])
        file_path = os.path.join(attachments_dir, file_name)

        # Scarica l'allegato
        with requests.get(f"{BASE_URL}{attachment_url}", auth=HTTPBasicAuth(API_USER, API_PASSWORD), stream=True) as r:
            r.raise_for_status()
            with open(file_path, "wb") as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)

# Funzione per scaricare le immagini e aggiornare i percorsi
def download_images(html_content, page_path):
    soup = BeautifulSoup(html_content, "html.parser")
    images = soup.find_all("img")
    for img in images:
        img_url = img.get("src")
        if img_url:
            img_name = os.path.basename(img_url)
            img_path = os.path.join(page_path, img_name)

            # Scarica l'immagine solo se non è già presente
            if not os.path.exists(img_path):
                img_data = requests.get(img_url).content
                with open(img_path, 'wb') as img_file:
                    img_file.write(img_data)

            # Aggiorna il percorso dell'immagine nell'HTML
            img['src'] = img_name
    return str(soup)

# Funzione per aggiungere il contenuto HTML nel DOCX
def add_html_to_docx(html_content, document):
    soup = BeautifulSoup(html_content, "html.parser")

    # Aggiungi il testo dei paragrafi
    for para in soup.find_all("p"):
        text = para.get_text()
        document.add_paragraph(text)

    # Aggiungi le tabelle
    for table in soup.find_all("table"):
        rows = table.find_all("tr")  # Trova tutte le righe della tabella
        # Calcola il numero massimo di celle in una riga
        max_cells = max(len(row.find_all("td")) for row in rows)

        # Crea la tabella con il numero massimo di celle
        table_doc = document.add_table(rows=0, cols=max_cells)

        # Aggiungi righe alla tabella
        for row in rows:
            row_cells = row.find_all("td")  # Trova tutte le celle della riga
            row_cells_text = [cell.get_text() for cell in row_cells]  # Estrai il testo da ogni cella
            
            # Aggiungi una nuova riga alla tabella
            table_row = table_doc.add_row().cells
            
            # Assegna il testo a ciascuna cella della riga
            for i, cell_text in enumerate(row_cells_text):
                table_row[i].text = cell_text  # Assegna il testo alla cella

            # Se ci sono celle vuote (ad esempio, se la riga ha meno celle di altre),
            # aggiungi celle vuote
            for i in range(len(row_cells_text), max_cells):
                table_row[i].text = ""  # Imposta il testo della cella vuota a ""
            
    # Aggiungi porzioni di codice (se presenti)
    for code_tag in soup.find_all("code"):
        code_text = code_tag.get_text()
        document.add_paragraph(code_text)  # Aggiungi il testo senza lo stile 'Code'


# Funzione per salvare il contenuto come DOCX
def save_as_docx(content, docx_path, page_path):
    # Prima scarica le immagini
    content = download_images(content, page_path)  # Scarica le immagini e aggiorna i percorsi

    document = Document()
    
    # Aggiungi il contenuto HTML al file DOCX
    add_html_to_docx(content, document)

    # Salva il file DOCX
    document.save(docx_path)

# Funzione principale per eseguire l'export
def export_space():
    pages = get_pages_with_hierarchy(SPACE_KEY)
    for page in pages:
        page_path = os.path.join(OUTPUT_DIR, get_hierarchy_path(page))
        os.makedirs(page_path, exist_ok=True)

        # Salva il contenuto della pagina come DOCX
        content = page["body"]["storage"]["value"]  # Contenuto HTML della pagina
        docx_path = os.path.join(page_path, f"{clean_name(page['title'])}.docx")
        save_as_docx(content, docx_path, page_path)

        # Scarica gli allegati
        download_attachments(page["id"], page_path)


# Esecuzione dello script
if __name__ == "__main__":
    export_space()
